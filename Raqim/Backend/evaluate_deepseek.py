import re
from pathlib import Path
from docx import Document
from jiwer import wer
from rapidfuzz.distance import Levenshtein
import pandas as pd


GROUND_TRUTH_DIR = "ground_truth"
DEEPSEEK_DIR = "deepseek_outputs"
OUTPUT_EXCEL = "deepseek_evaluation.xlsx"


def normalize_text(text):
    text = text.replace("\u200f", "").replace("\u200e", "")
    text = text.replace("ـ", "")
    text = re.sub(r"[ًٌٍَُِّْ]", "", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def read_docx(path):
    doc = Document(path)
    parts = []

    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return normalize_text("\n".join(parts))


def evaluate_pair(gt_path, pred_path):
    gt = read_docx(gt_path)
    pred = read_docx(pred_path)

    cer = Levenshtein.distance(gt, pred) / max(len(gt), 1)
    wer_score = wer(gt, pred)

    return {
        "file": gt_path.name,
        "CER_percent": cer * 100,
        "WER_percent": wer_score * 100,
        "Character_Accuracy": 100 - (cer * 100),
        "Word_Accuracy": 100 - (wer_score * 100),
        "GT_chars": len(gt),
        "DeepSeek_chars": len(pred),
    }


def main():
    gt_dir = Path(GROUND_TRUTH_DIR)
    pred_dir = Path(DEEPSEEK_DIR)

    results = []

    for gt_file in gt_dir.glob("*.docx"):
        pred_file = pred_dir / gt_file.name

        if not pred_file.exists():
            print(f"Missing DeepSeek output for: {gt_file.name}")
            continue

        print(f"Evaluating: {gt_file.name}")
        results.append(evaluate_pair(gt_file, pred_file))

    df = pd.DataFrame(results)

    summary = pd.DataFrame([{
        "Average_CER_percent": df["CER_percent"].mean(),
        "Average_WER_percent": df["WER_percent"].mean(),
        "Average_Character_Accuracy": df["Character_Accuracy"].mean(),
        "Average_Word_Accuracy": df["Word_Accuracy"].mean(),
        "Files_Count": len(df)
    }])

    with pd.ExcelWriter(OUTPUT_EXCEL, engine="openpyxl") as writer:
        df.to_excel(writer, sheet_name="details", index=False)
        summary.to_excel(writer, sheet_name="summary", index=False)

    print("Done.")
    print(summary)
    print(f"Saved to: {OUTPUT_EXCEL}")


if __name__ == "__main__":
    main()