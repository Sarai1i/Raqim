"""Shared DOCX builder for Raqim.

This module reconstructs a Word (.docx) document from Raqim's OCR page/word
structure (the list of word dicts tagged with ``block_id`` / ``block_type`` /
``table_*`` / ``line_index`` / ``is_math`` ...). It is engine-agnostic: any OCR
engine (DeepSeek-OCR-2 or DotOCR) that emits this structure gets the same layout
reconstruction:

* ``block_type == "table"``   -> a real Word table (with merged cells, RTL).
* ``block_type == "heading"`` -> a Word heading (level 2).
* otherwise                    -> RTL paragraphs in reading order, with inline
                                  math rendered as OMML equations.

It was extracted out of ``app.py`` so the layout reconstruction can be reused by
the OCR engines and unit-tested without importing the Flask app.
"""

from __future__ import annotations

import re


def _word_display(word_data):
    """النص المعروض للكلمة مع تفضيل التصحيح اليدوي."""
    return (word_data.get("corrected_word") or word_data.get("word") or "").strip()


def _set_rtl_paragraph(paragraph):
    """ضبط اتجاه الفقرة من اليمين لليسار (للعربية)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)


def _fix_omml_bugs(omml):
    """يصلّح أخطاء معروفة في ناتج mathml2omml.

    أبرزها: عند تحويل ``\\vec`` (سهم فوق الحرف) تُغلق المكتبة الوسم خطأً بـ
    ``</m:groupChr>`` بدل ``</m:groupChrPr>``، فينكسر XML. نصلّح الإغلاق حتى
    يبقى السهم الحقيقي فوق المتجه بدل اللجوء إلى بديل.
    """
    omml = omml.replace('<m:pos m:val="top"/></m:groupChr>', '<m:pos m:val="top"/></m:groupChrPr>')
    return omml


def _tidy_equation_latex(latex):
    r"""تنظيف أخطاء OCR الشائعة داخل معادلات LaTeX قبل تحويلها.

    مثال: يقرأ DeepSeek الصفر المنخفض في ``\mu_0`` أحيانًا كفاصلة (``\mu_,``)
    أو ``\mu_{,}``، فنعيدها إلى صفر. كما نزيل فراغات زائدة.
    """
    if not latex:
        return latex
    latex = re.sub(r"_\{?\s*[,،]\s*\}?", "_{0}", latex)  # منخفض = فاصلة -> صفر
    latex = re.sub(r"\\bullet\s*", "", latex)  # إزالة نقطة تعداد تسرّبت للمعادلة
    return latex.strip()


def _latex_to_omml_element(latex):
    """حوّل LaTeX إلى عنصر معادلة Word حقيقية (OMML). يرجع None عند الفشل."""
    try:
        import latex2mathml.converter as _l2m
        import mathml2omml
        from docx.oxml import parse_xml
        mathml = _l2m.convert(_tidy_equation_latex(latex))
        omml = _fix_omml_bugs(mathml2omml.convert(mathml))
        omml = omml.replace(
            "<m:oMath>",
            '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">',
            1,
        )
        return parse_xml(omml)
    except Exception as e:
        print(f"⚠️ تعذّر تحويل معادلة إلى Word، سيُكتب نصها كما هو: {str(e)[:80]}")
        return None


def _add_math_paragraph(document, latex, align_right, _set_rtl_paragraph):
    """يضيف فقرة تحوي معادلة Word حقيقية، أو نص LaTeX احتياطيًا عند فشل التحويل."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = document.add_paragraph()
    element = _latex_to_omml_element(latex)
    if element is not None:
        p._p.append(element)
    else:
        run = p.add_run(f"$ {latex} $")  # احتياطي: نص LaTeX واضح
        run.font.name = "Cambria Math"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # المعادلات تُوسّط عادةً
    return p


def build_corrected_docx(path, pages):
    """بناء ملف Word (.docx) من نتائج OCR: فقرات وعناوين وجداول حقيقية مع RTL.

    يعتمد على وسوم block_id/block_type/table_* التي يضيفها محرك OCR، فيعيد بناء
    البنية: العناوين كعناوين، الفقرات كفقرات، والجداول كجداول Word بخلايا حقيقية.
    ``pages`` هي قائمة الصفحات (كل صفحة تحوي ``text`` بقائمة الكلمات الموسومة).
    """
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    source_pages = pages or []

    document = Document()
    # خط افتراضي يدعم العربية
    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(12)

    for page in source_pages:
        words = page.get("text", [])
        if not words:
            continue

        # تجميع الكلمات إلى كتل بالترتيب حسب block_id
        blocks = []
        for word in words:
            bid = word.get("block_id")
            btype = word.get("block_type", "text")
            if blocks and blocks[-1]["id"] == bid and blocks[-1]["type"] == btype:
                blocks[-1]["words"].append(word)
            else:
                blocks.append({"id": bid, "type": btype, "words": [word]})

        for block in blocks:
            if block["type"] == "table":
                # إعادة بناء الجدول مع دعم الخلايا المدموجة (rowspan/colspan).
                cells_info = []
                nrows = 0
                ncols = 0
                for w in block["words"]:
                    r = int(w.get("table_row", 0))
                    c = int(w.get("table_col", 0))
                    rspan = max(1, int(w.get("table_rowspan", 1)))
                    cspan = max(1, int(w.get("table_colspan", 1)))
                    cells_info.append((r, c, rspan, cspan, _word_display(w)))
                    nrows = max(nrows, r + rspan)
                    ncols = max(ncols, c + cspan)
                if not cells_info or nrows < 1 or ncols < 1:
                    continue

                table = document.add_table(rows=nrows, cols=ncols)
                table.style = "Table Grid"
                table.alignment = WD_TABLE_ALIGNMENT.RIGHT
                # اجعل الجدول من اليمين لليسار: العمود 0 يظهر على اليمين (ترتيب عربي).
                from docx.oxml.ns import qn as _qn
                from docx.oxml import OxmlElement as _OxmlElement
                _tblPr = table._tbl.tblPr
                _bidi = _OxmlElement("w:bidiVisual")
                _tblPr.append(_bidi)

                for (r, c, rspan, cspan, text) in cells_info:
                    top_left = table.cell(r, c)
                    # دمج الخلايا إن كان هناك امتداد
                    if rspan > 1 or cspan > 1:
                        bottom_right = table.cell(min(r + rspan - 1, nrows - 1),
                                                  min(c + cspan - 1, ncols - 1))
                        try:
                            merged = top_left.merge(bottom_right)
                        except Exception:
                            merged = top_left
                        target = merged
                    else:
                        target = top_left
                    target.text = text
                    for p in target.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        _set_rtl_paragraph(p)
                document.add_paragraph("")  # مسافة بعد الجدول
            elif block["type"] == "heading":
                heading_text = " ".join(
                    _word_display(w) for w in block["words"] if _word_display(w)
                ).strip()
                if heading_text:
                    p = document.add_heading(level=2)
                    run = p.add_run(heading_text)
                    run.font.name = "Arial"
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    _set_rtl_paragraph(p)
            else:
                # نبني الكتلة سطرًا بسطر مع دعم المعادلات داخل السطر (inline).
                # كل سطر يحوي كلماته بالترتيب (نص أو معادلة)، فنضيفها في فقرة واحدة.
                from docx.oxml.ns import qn

                lines_words = {}
                for w in block["words"]:
                    li = w.get("line_index", 0)
                    lines_words.setdefault(li, []).append(w)
                line_ids = sorted(lines_words.keys())
                if not line_ids:
                    continue

                # هل السطر كله معادلة معروضة وحيدة؟ (لتوسيطها لوحدها)
                def _line_is_single_display(line_ws):
                    real = [w for w in line_ws if _word_display(w) or w.get("is_math")]
                    return len(real) == 1 and real[0].get("is_math") and real[0].get("math_display")

                for li in line_ids:
                    line_ws = lines_words[li]

                    if _line_is_single_display(line_ws):
                        mw = next(w for w in line_ws if w.get("is_math"))
                        _add_math_paragraph(document, mw.get("latex") or mw.get("word"), True, _set_rtl_paragraph)
                        continue

                    # فقرة عادية: نضيف كل كلمة كنص أو معادلة inline بالترتيب.
                    p = document.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    _set_rtl_paragraph(p)
                    buffer = []

                    def _flush_buffer(par):
                        if buffer:
                            text = " ".join(buffer).strip()
                            # تنظيف المسافة بعد نقطة التعداد: "•M" -> "• M"
                            text = re.sub(r"^([•\-\*])\s*", r"\1 ", text)
                            if text:
                                run = par.add_run(text)
                                run.font.name = "Arial"
                            buffer.clear()

                    n_items = len(line_ws)
                    for pos, w in enumerate(line_ws):
                        if w.get("is_math"):
                            _flush_buffer(p)
                            element = _latex_to_omml_element(w.get("latex") or w.get("word"))
                            # مسافة قبل المعادلة إن سبقها نص.
                            if pos > 0:
                                p.add_run(" ")
                            if element is not None:
                                p._p.append(element)
                            else:
                                run = p.add_run(f"$ {w.get('latex')} $")
                                run.font.name = "Cambria Math"
                            # مسافة بعد المعادلة إن تبعها نص.
                            if pos < n_items - 1 and not line_ws[pos + 1].get("is_math"):
                                p.add_run(" ")
                        else:
                            disp = _word_display(w)
                            if disp:
                                buffer.append(disp)
                    _flush_buffer(p)
                continue

        document.add_page_break()

    document.save(path)
